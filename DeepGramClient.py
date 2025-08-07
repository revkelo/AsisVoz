import os
import time
from deepgram import DeepgramClient, PrerecordedOptions
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import locale


class DeepgramPDFTranscriber:
    def __init__(self, api_key: str):
        if not api_key:
            raise ValueError("API key de Deepgram requerida.")
        print("[Inicializando] Cliente Deepgram...")
        self.client = DeepgramClient(api_key)

    def segundos_a_hhmmss(self, segundos: float) -> str:
        horas = int(segundos // 3600)
        minutos = int((segundos % 3600) // 60)
        segundos_restantes = int(segundos % 60)
        return f"[{horas:02d}:{minutos:02d}:{segundos_restantes:02d}]"

    def generar_pdf(self, nombre_salida: str, transcripciones: list[str]):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, "Transcripción de audio", ln=True, align="C")
        pdf.set_font("Arial", size=10)
        pdf.cell(0, 10, f"Archivo: {nombre_salida}", ln=True, align="C")
        pdf.ln(10)

        pdf.set_font("Arial", size=12)
        for linea in transcripciones:
            pdf.multi_cell(0, 10, linea)

        if not nombre_salida.lower().endswith(".pdf"):
            nombre_salida += ".pdf"

        pdf.output(nombre_salida)
        print(f"\n✅ PDF guardado como '{nombre_salida}'")

    def transcribir_audio(self, ruta_audio, nombre_salida):
        inicio = time.time()
        try:
            if not ruta_audio or not os.path.isfile(ruta_audio):
                raise FileNotFoundError("❌ Archivo no válido o no encontrado.")

            with open(ruta_audio, "rb") as f:
                audio_bytes = f.read()

            options = PrerecordedOptions(
                model="nova-2",
                language="es",
                smart_format=True,
                punctuate=True,
                paragraphs=True,
                diarize=True,
            )

            response = self.client.listen.prerecorded.v("1").transcribe_file(
                {"buffer": audio_bytes},
                options,
                timeout=1000
            )

            channel = response.results.channels[0]
            transcripciones = []

            for alt in channel.alternatives:
                if hasattr(alt, "paragraphs") and alt.paragraphs and hasattr(alt.paragraphs, "paragraphs"):
                    for paragraph in alt.paragraphs.paragraphs:
                        tiempo = self.segundos_a_hhmmss(paragraph.start)
                        speaker = f"Locutor {paragraph.speaker}"
                        texto = " ".join([s.text for s in paragraph.sentences])
                        linea = f"{tiempo} {speaker}: {texto.strip()}"
                        transcripciones.append(linea)
                elif alt.transcript and alt.transcript.strip():
                    # Fallback si hay texto sin párrafos
                    transcripciones.append(alt.transcript.strip())

            if not transcripciones:
                raise ValueError("⚠ No se detectó voz en el archivo. Verifica que contenga audio hablado.")

            #self.generar_pdf(nombre_salida, transcripciones)
            self.generar_word(nombre_salida, transcripciones, ruta_audio)





            fin = time.time()
            print(f"\n⏱ Tiempo de ejecución: {fin - inicio:.2f} segundos")

        except Exception as e:
            print(f"❌ Exception: {e}")
            raise  # Opcional: relanza para que la GUI también lo muestre si lo capturas desde allá

    def obtener_fecha_creacion(self, ruta_archivo):
        try:
            # En Windows, os.path.getctime da la fecha de creación
            fecha_creacion = os.path.getctime(ruta_archivo)
            fecha_str = f" Fecha de creación: {time.strftime('%Y-%m-%d %I:%M:%S %p', time.localtime(fecha_creacion))}"
            return fecha_str
        except Exception as e:
            return f" Fecha de creación (modificación): Error al obtener la fecha: {e}"


    def generar_word(self, nombre_salida: str, transcripciones: list[str], ruta_audio: str):
        doc = Document()

        # Título centrado y en negrita
        titulo = doc.add_heading("Transcripción de audio", level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Nombre del archivo (solo el nombre, no toda la ruta)
        nombre_archivo = os.path.basename(ruta_audio)
        archivo_parrafo = doc.add_paragraph(f"Archivo: {nombre_archivo}")
        archivo_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        archivo_parrafo.runs[0].font.size = Pt(10)

        fecha_formateada = self.obtener_fecha_creacion(ruta_audio);

        fecha_parrafo = doc.add_paragraph(f"Fecha de creación: {fecha_formateada}")
        fecha_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_parrafo.runs[0].font.size = Pt(10)

        # Espacio
        doc.add_paragraph("")

        # Texto de la transcripción
        for linea in transcripciones:
            parrafo = doc.add_paragraph()
            run = parrafo.add_run(linea)
            run.font.size = Pt(12)

        if not nombre_salida.lower().endswith(".docx"):
            nombre_salida += ".docx"

        doc.save(nombre_salida)
        print(f"\n✅ Word guardado como '{nombre_salida}'")

